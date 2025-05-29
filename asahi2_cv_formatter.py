# Professional Asahi CV Formatter - Clean & Simple Design
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import fitz  # PyMuPDF
import re
from PIL import Image
from collections import defaultdict
import time

# --- Professional Clean CSS Design ---
def apply_professional_css():
    st.markdown("""
    <style>
    .main-header {
        background: linear-gradient(90deg, #0f2027 0%, #2c5364 100%);
        color: #fff;
        padding: 1.5rem 1rem 1rem 1rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        text-align: center;
        font-size: 2.1rem;
        font-weight: 700;
        letter-spacing: 1px;
        box-shadow: 0 2px 8px rgba(44,83,100,0.12);
        cursor: pointer;
        position: relative;
        transition: all 0.3s ease;
    }
    .main-header:hover {
        background: linear-gradient(90deg, #1a2b32 0%, #3a6374 100%);
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(44,83,100,0.2);
    }
    .main-header .link-symbol {
        opacity: 0;
        transition: opacity 0.3s ease;
        margin-left: 0.5rem;
        font-size: 1.2rem;
        vertical-align: middle;
    }
    .main-header:hover .link-symbol {
        opacity: 1;
    }
    .main-header a {
        color: white;
        text-decoration: none;
        display: block;
        width: 100%;
        height: 100%;
    }
    .main-header a:hover {
        color: white;
        text-decoration: none;
    }
    .emoji {
        font-size: 2.2rem;
        vertical-align: middle;
        margin-right: 0.5rem;
    }
    .supported-formats {
        background: linear-gradient(135deg, #f0f4f8 0%, #e8f2ff 100%);
        padding: 1.5rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        border-left: 4px solid #2c5364;
        box-shadow: 0 2px 10px rgba(44,83,100,0.08);
    }
    .supported-formats h3 {
        color: #2c5364;
        margin-bottom: 1rem;
        font-weight: 600;
        font-size: 1.1rem;
    }
    .format-list {
        margin: 0;
        padding-left: 1.2rem;
        color: #4a5568;
    }
    .format-list li {
        margin-bottom: 0.5rem;
        font-weight: 500;
    }
    .pdf-icon {
        color: #e25555;
        font-size: 1.1rem;
        margin-left: 0.5rem;
    }
    .docx-icon {
        color: #4a90e2;
        font-size: 1.1rem;
        margin-left: 0.5rem;
    }
    .clean-card {
        background: #ffffff;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.04);
        border: 1px solid #e2e8f0;
        margin-bottom: 1.5rem;
    }
    .clean-card h3 {
        color: #374151;
        font-weight: 600;
        margin-bottom: 1.5rem;
        font-size: 1.2rem;
        border-bottom: 2px solid #f1f5f9;
        padding-bottom: 0.5rem;
    }
    .status-success {
        color: #2d5016;
        padding: 0.5rem 0;
        margin: 1rem 0;
        font-weight: 500;
        font-size: 0.95rem;
    }
    .status-info {
        background: linear-gradient(135deg, #4299e1 0%, #3182ce 100%);
        color: white;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        font-weight: 500;
        box-shadow: 0 2px 8px rgba(66,153,225,0.2);
    }
    .status-warning {
        background: linear-gradient(135deg, #ed8936 0%, #dd6b20 100%);
        color: white;
        padding: 1rem 1.5rem;
        border-radius: 10px;
        margin: 1rem 0;
        font-weight: 500;
        box-shadow: 0 2px 8px rgba(237,137,54,0.2);
    }
    .progress-step {
        margin: 0.5rem 0;
        color: #4a5568;
        font-weight: 500;
    }
    .step-check {
        color: #48bb78;
        margin-right: 0.5rem;
        font-weight: bold;
    }
    .metric-item {
        text-align: center;
        background: #f7fafc;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #e2e8f0;
    }
    .metric-number {
        font-size: 1.8rem;
        font-weight: bold;
        color: #2c5364;
    }
    .metric-label {
        color: #718096;
        font-size: 0.9rem;
        margin-top: 0.2rem;
    }
    .detection-summary {
        background: #fffbeb;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #f59e0b;
    }
    .detection-item {
        margin: 0.5rem 0;
        color: #92400e;
        font-size: 0.9rem;
    }
    .footer {
        margin-top: 3rem;
        padding-top: 2rem;
        border-top: 1px solid #e2e8f0;
        color: #718096;
        font-size: 0.9rem;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)

# --- Advanced PII Detection Class ---
class PIIDetector:
    def __init__(self):
        self.patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}|\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),
            'address': re.compile(r'\d+\s+[\w\s,.-]+(?:street|st|avenue|ave|road|rd|drive|dr|lane|ln|boulevard|blvd|court|ct|place|pl)(?:\s+(?:apt|apartment|unit|#)\s*\w+)?', re.IGNORECASE),
            'zip_code': re.compile(r'\b\d{5}(?:-\d{4})?\b'),
            'height': re.compile(r'\b(?:\d+\'\s*\d+\"|\d+\s*ft\s*\d+\s*in|\d+\.\d+\s*m|\d+\s*cm)\b', re.IGNORECASE),
            'weight': re.compile(r'\b\d+(?:\.\d+)?\s*(?:lbs?|pounds?|kg|kilograms?)\b', re.IGNORECASE),
            'date_of_birth': re.compile(r'\b(?:DOB|Date of Birth|Born):?\s*(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{2,4})', re.IGNORECASE),
            'ssn': re.compile(r'\b\d{3}-\d{2}-\d{4}\b'),
            'linkedin': re.compile(r'linkedin\.com/in/[\w-]+', re.IGNORECASE),
	    'name': re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+$'),

        }
        
        self.personal_keywords = [
            'home address', 'residential address', 'current address', 'permanent address',
            'contact number', 'mobile number', 'cell phone', 'telephone', 'tel. no', 'phone',
            'date of birth', 'dob', 'born on', 'age:', 'years old',
            'marital status', 'married', 'single', 'divorced',
            'nationality', 'citizen', 'passport', 'visa status',
            'height:', 'weight:', 'blood type', 'emergency contact',
            'email address', 'e-mail', 'gmail', 'yahoo','name','full name', '@'
        ]
        
        # Patterns to identify lines containing personal information that should be completely removed
        self.pii_line_patterns = [
            re.compile(r'.*(?:tel\.?\s*no\.?|phone|mobile|contact).*?[\+\(]?\d{1,4}[\s\-\(\)]*\d{3,4}[\s\-]*\d{3,4}.*', re.IGNORECASE),
            re.compile(r'.*email.*?[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}.*', re.IGNORECASE),
            re.compile(r'.*(?:address|location).*?\d+.*?(?:street|st|avenue|ave|road|rd|drive|dr|lane|ln|boulevard|blvd).*', re.IGNORECASE),
            re.compile(r'.*(?:height|weight|born|dob|date of birth).*', re.IGNORECASE),
            re.compile(r'.*(?:nationality|citizenship|passport|visa).*', re.IGNORECASE),
            re.compile(r'.*(?:marital|married|single|divorced).*', re.IGNORECASE),
        ]
    
    def detect_names(self, text):
        detected_names = set()
        name_patterns = [
            # Regular capitalized names (John Doe)
            re.compile(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]*)*)\s*$', re.MULTILINE),
            # Names with labels
            re.compile(r'(?:Name|Full Name|Candidate):?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]*)+)', re.IGNORECASE),
            # Names at start of line
            re.compile(r'^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', re.MULTILINE),
            # ALL CAPS NAMES (NEW PATTERN)
            re.compile(r'^([A-Z]{2,}(?:\s+[A-Z]{2,})+)\s*$', re.MULTILINE),
            # Mixed case ALL CAPS names
            re.compile(r'(?:Name|Full Name|Candidate):?\s*([A-Z]{2,}(?:\s+[A-Z]{2,})+)', re.IGNORECASE),
            # ALL CAPS at start of line
            re.compile(r'^([A-Z]{2,}\s+[A-Z]{2,}(?:\s+[A-Z]{2,})*)', re.MULTILINE),
        ]
        
        for pattern in name_patterns:
            matches = pattern.findall(text)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0] if match[0] else match[1]
                if not any(word.lower() in ['university', 'college', 'company', 'corporation', 'inc', 'ltd', 'experience', 'education', 'skills', 'objective', 'summary', 'profile', 'references', 'qualifications'] for word in match.split()):
                    # For ALL CAPS, ensure it's at least 2 words and each word is at least 2 characters
                    if match.isupper():
                        words = match.split()
                        if len(words) >= 2 and all(len(word) >= 2 for word in words):
                            detected_names.add(match.strip())
                    # For regular names, ensure at least 2 words
                    elif len(match.split()) >= 2:
                        detected_names.add(match.strip())
        
        return list(detected_names)
    
    def detect_all_pii(self, text):
        detected_pii = defaultdict(list)
        # Don't include names in the returned PII
        
        for pii_type, pattern in self.patterns.items():
            matches = pattern.findall(text)
            if matches:
                detected_pii[pii_type] = list(set(matches))
        
        personal_info_lines = []
        lines = text.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in self.personal_keywords):
                personal_info_lines.append(line.strip())
        
        if personal_info_lines:
            detected_pii['personal_info_lines'] = personal_info_lines
        
        return dict(detected_pii)
    
    def remove_pii(self, text, detected_pii):
        cleaned_text = text
        removal_count = 0
        
        # Still detect and remove names internally, but don't show them in PII report
        detected_names = self.detect_names(text)
        for name in detected_names:
            if name and len(name.strip()) > 2:
                # Use word boundary matching for better accuracy
                pattern = re.compile(r'\b' + re.escape(name) + r'\b', re.IGNORECASE)
                if pattern.search(cleaned_text):
                    cleaned_text = pattern.sub('', cleaned_text)
                    removal_count += 1
        
        # Process line by line to completely remove PII-containing lines
        lines = cleaned_text.split('\n')
        filtered_lines = []
        
        for line in lines:
            line_contains_pii = False
            original_line = line.strip()
            
            # Check if line matches any PII line patterns (complete removal)
            for pii_pattern in self.pii_line_patterns:
                if pii_pattern.match(line):
                    line_contains_pii = True
                    removal_count += 1
                    break
            
            # Check for individual PII items in the line
            if not line_contains_pii:
                for pii_type, items in detected_pii.items():
                    for item in items:
                        if item and len(str(item).strip()) > 1:
                            if str(item).lower() in line.lower():
                                line_contains_pii = True
                                removal_count += 1
                                break
                    if line_contains_pii:
                        break
            
            # Check for personal keywords that indicate the entire line should be removed
            if not line_contains_pii:
                line_lower = line.lower()
                for keyword in self.personal_keywords:
                    if keyword in line_lower:
                        # Additional check to avoid removing work-related lines
                        if not any(work_keyword in line_lower for work_keyword in ['experience', 'work', 'employment', 'company', 'project', 'skill', 'education', 'university', 'college']):
                            line_contains_pii = True
                            removal_count += 1
                            break
            
            # Only keep lines that don't contain PII
            if not line_contains_pii and original_line:
                filtered_lines.append(original_line)
        
        cleaned_text = '\n'.join(filtered_lines)
        
        # Clean up extra whitespace and empty lines
        cleaned_text = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_text)
        cleaned_text = cleaned_text.strip()
        
        return cleaned_text, removal_count

# --- Helper Functions ---
def extract_text_from_pdf(file):
    text = ""
    try:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""
    return text

def extract_text_from_docx(file):
    try:
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def abbreviate_name_age(full_name, age):
    try:
        name_parts = [part.strip() for part in full_name.strip().split() if part.strip()]
        if not name_parts:
            return f"N.A.{age}yrs"
        initials = ''.join([part[0].upper() + '.' for part in name_parts])
        return f"{initials} {age}yrs"
    except Exception:
        return f"N.A.{age}yrs"

def add_header_with_logo(doc, logo_img):
    section = doc.sections[0]
    header = section.header
    
    for paragraph in header.paragraphs:
        paragraph.clear()
    
    logo_para = header.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    tab_stops = logo_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    
    logo_run = logo_para.add_run("\t")
    
    image_stream = BytesIO()
    logo_img.save(image_stream, format='PNG')
    image_stream.seek(0)
    logo_run.add_picture(image_stream, width=Inches(2.634), height=Inches(0.508))
    
    section.header_distance = Inches(0.4)

def generate_asahi_cv(cleaned_text, logo_img, candidate_name, age):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    add_header_with_logo(doc, logo_img)
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    name_paragraph = doc.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_paragraph.paragraph_format.space_after = Pt(24)
    
    name_run = name_paragraph.add_run(abbreviate_name_age(candidate_name, age))
    name_run.font.name = 'ＭＳ 明朝'
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    
    doc.add_paragraph()
    
    content_lines = [line.strip() for line in cleaned_text.strip().split("\n") if line.strip()]
    for line in content_lines:
        if line.strip():
            doc.add_paragraph(line.strip())
    
    return doc

# --- Main Application ---
def main():
    st.set_page_config(
        page_title="Asahi CV Formatter",
        layout="centered",
        initial_sidebar_state="collapsed"
    )
    
    apply_professional_css()
    
    # Clickable header with hover link symbol effect - FIXED
    st.markdown("""
    <div class="main-header" onclick="document.getElementById('upload-section').scrollIntoView({behavior: 'smooth'});">
        <span class="emoji">📝</span>Asahi CV Formatter<span class="link-symbol">🔗</span>
    </div>
    """, unsafe_allow_html=True)
    st.write("Professional CV formatting with automatic privacy protection")

    # Supported formats section with improved design
    st.markdown("""
    <div class="supported-formats">
        <h3>📄 Supported Formats</h3>
        <ul class="format-list">
            <li>PDF documents <span class="pdf-icon">📄</span></li>
            <li>DOCX documents <span class="docx-icon">📄</span></li>
        </ul>
    </div>
    """, unsafe_allow_html=True)
    
    pii_detector = PIIDetector()
    
    # Upload section - Clean version without extra spacing
    st.markdown('<div id="upload-section"></div>', unsafe_allow_html=True)
    uploaded_file = st.file_uploader(
        "📄 Choose CV file (PDF or DOCX)", 
        type=["docx", "pdf"],
        help="Upload the candidate's resume in PDF or Word format"
    )
    
    # Candidate information inputs
    st.markdown("### 👤 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        candidate_name = st.text_input("👤 Candidate Full Name", placeholder="e.g., John Doe Smith")
    with col2:
        age = st.number_input("🎂 Candidate Age", min_value=18, max_value=99, step=1)
    
    # Processing section - Auto-process when file, name and age are provided
    if uploaded_file and candidate_name.strip() and age:
        # Load logo
        try:
            logo_img = Image.open("asahi_logo-04.jpg")
        except FileNotFoundError:
            st.markdown("""
            <div class="status-warning">
                <strong>Warning:</strong> Logo file 'asahi_logo-04.jpg' not found. Please ensure it's in the same directory.
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        # Extract text
        if uploaded_file.name.lower().endswith(".pdf"):
            raw_text = extract_text_from_pdf(uploaded_file)
        else:
            raw_text = extract_text_from_docx(uploaded_file)
        
        if not raw_text.strip():
            st.markdown("""
            <div class="status-warning">
                <strong>Error:</strong> No text could be extracted from the file. Please check the file format.
            </div>
            """, unsafe_allow_html=True)
            st.stop()
        
        # Show file loaded successfully in plain text
        st.markdown(f"""
        <div class="status-success">
            File loaded successfully: {uploaded_file.name} ({len(raw_text.split())} words)
        </div>
        """, unsafe_allow_html=True)
        
        # Auto-process without button
        with st.spinner("Processing CV..."):
            # Use the manually entered candidate name
            
            # Detect and remove ALL PII including names
            detected_pii = pii_detector.detect_all_pii(raw_text)
            cleaned_text, removal_count = pii_detector.remove_pii(raw_text, detected_pii)
            
            # Generate document with only abbreviation in header
            final_doc = generate_asahi_cv(cleaned_text, logo_img, candidate_name, age)
            buffer = BytesIO()
            final_doc.save(buffer)
            buffer.seek(0)
            
            # Show simple completion message
            st.markdown("""
            <div class="status-success">
                CV Processing Complete!
            </div>
            """, unsafe_allow_html=True)
            
            # Download with abbreviation filename
            abbreviation = abbreviate_name_age(candidate_name, age).replace(f" {age}yrs", "")
            file_name = f"Asahi_CV_{abbreviation}.docx"
            st.download_button(
                label="Download Formatted CV",
                data=buffer,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    elif uploaded_file or candidate_name.strip() or age:
        st.markdown("""
        <div class="status-info">
            <strong>Ready to process:</strong> Please provide all required information above to continue.
        </div>
        """, unsafe_allow_html=True)

    # Footer
    st.markdown('<div class="footer">©Asahi Kogyo Co., Ltd. Osaka Office</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
